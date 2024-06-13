import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
from docx import Document
from docx.shared import Inches
import os

def load_data(filename):
    """ Load data from a CSV file """
    return pd.read_csv(filename)

def calculate_kpis(df, doc):
    """ Calculate and print KPIs for YouTube videos, add to doc """
    kpis = [
        ("Total Videos", df.shape[0]),
        ("Total Views", df['views'].sum()),
        ("Total Likes", df['likes'].sum()),
        ("Total Comments", df['comments'].sum()),
        ("Average Views per Video", df['views'].mean()),
        ("Average Likes per Video", df['likes'].mean()),
        ("Average Comments per Video", df['comments'].mean()),
        ("Most Views", df['views'].max()),
        ("Top 5% Average Views", df[df['views'] >= df['views'].quantile(0.95)]['views'].mean()),
        ("Average Video Age (in days)", (datetime.now().replace(tzinfo=None) - pd.to_datetime(df['published_at']).dt.tz_localize(None)).dt.days.mean()),
        ("Average Likes per View", df['likes'].sum() / df['views'].sum()),
        ("Average Comments per View", df['comments'].sum() / df['views'].sum()),
        ("Engagement Rate", (df['likes'].sum() + df['comments'].sum()) / df['views'].sum())
    ]
    doc.add_heading('Key Performance Indicators', level=1)
    for kpi_name, kpi_value in kpis:
        doc.add_paragraph(f"{kpi_name}: {kpi_value:.2f}" if isinstance(kpi_value, float) else f"{kpi_name}: {kpi_value}")

def plot_and_save(df, filename):
    """ Plot and save the figure to include in Word doc """
    plt.figure(figsize=(10, 5))
    df['views'].plot(title='Views Over Time')
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

def add_chart_to_doc(doc, filename):
    """ Add chart image to Word document """
    doc.add_picture(filename, width=Inches(6.0))

def main():
    filename = 'videos.csv'  # Path to the CSV file
    df = load_data(filename)
    
    # Convert columns to numeric if they are not already
    df[['views', 'likes', 'comments']] = df[['views', 'likes', 'comments']].apply(pd.to_numeric, errors='coerce')
    
    doc = Document()
    calculate_kpis(df, doc)
    
    # Save plot and add to document
    plot_filename = 'views_over_time.png'
    plot_and_save(df.set_index(pd.to_datetime(df['published_at']).dt.tz_localize(None)), plot_filename)
    add_chart_to_doc(doc, plot_filename)
    
    # Save the document
    doc.save('YouTube_Analytics_Report.docx')
    # Clean up plot file
    os.remove(plot_filename)

if __name__ == '__main__':
    main()
